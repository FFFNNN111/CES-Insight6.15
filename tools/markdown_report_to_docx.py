from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path("literature_reports/2026-06-01_literature_scan.md")
OUTPUT = Path("literature_reports/2026-06-01_literature_scan.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def normalize_inline(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    return text


def add_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(normalize_inline(part[2:-2]))
            run.bold = True
        else:
            paragraph.add_run(normalize_inline(part))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.space_after = Pt(4)

    for style_name, size, color in [
        ("Title", 22, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 11.5, "404040"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    quote_style = styles["Intense Quote"]
    quote_style.font.name = "Microsoft YaHei"
    quote_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    quote_style.font.size = Pt(9.5)
    quote_style.font.color.rgb = RGBColor.from_string("404040")
    quote_style.paragraph_format.left_indent = Cm(0.55)
    quote_style.paragraph_format.right_indent = Cm(0.3)


def split_table_row(line: str) -> list[str]:
    body = line.strip().strip("|")
    return [normalize_inline(cell.strip()) for cell in body.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (col_count - len(row)))

    table = document.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.allow_autofit = False

    usable_width = int(15120)
    if col_count >= 9:
        weights = [1.8, 0.7, 1.2, 1.5, 1.0, 1.2, 1.0, 2.0, 1.6, 1.0, 0.9]
        weights = weights[:col_count]
    else:
        weights = [1.0] * col_count
    total_weight = sum(weights)
    widths = [int(usable_width * weight / total_weight) for weight in weights]

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cell = cells[c_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(7.5 if col_count >= 9 else 9)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("FFFFFF")
                set_cell_shading(cell, "2F5597")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_width(cell, widths[c_idx])
    set_repeat_table_header(table.rows[0])

    document.add_paragraph()


def add_paragraph(document: Document, line: str) -> None:
    if not line.strip():
        return

    if line.startswith("# "):
        paragraph = document.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(paragraph, line[2:].strip())
        return

    heading_match = re.match(r"^(#{2,6})\s+(.+)$", line)
    if heading_match:
        level = min(len(heading_match.group(1)) - 1, 3)
        paragraph = document.add_paragraph(style=f"Heading {level}")
        add_runs(paragraph, heading_match.group(2).strip())
        return

    if line.startswith("- "):
        paragraph = document.add_paragraph(style="List Bullet")
        add_runs(paragraph, line[2:].strip())
        return

    number_match = re.match(r"^(\d+)\.\s+(.+)$", line)
    if number_match:
        paragraph = document.add_paragraph(style="List Number")
        add_runs(paragraph, number_match.group(2).strip())
        return

    if line.startswith("> "):
        paragraph = document.add_paragraph(style="Intense Quote")
        add_runs(paragraph, line[2:].strip())
        return

    paragraph = document.add_paragraph()
    add_runs(paragraph, line.strip())


def build_docx(source: Path, output: Path) -> None:
    document = Document()
    configure_document(document)

    lines = source.read_text(encoding="utf-8").splitlines()
    table_rows: list[list[str]] = []
    in_table = False

    for raw_line in lines:
        line = raw_line.rstrip()
        if "|" in line and line.strip().startswith("|"):
            if is_table_separator(line):
                in_table = True
                continue
            table_rows.append(split_table_row(line))
            in_table = True
            continue

        if in_table:
            add_table(document, table_rows)
            table_rows = []
            in_table = False

        add_paragraph(document, line)

    if table_rows:
        add_table(document, table_rows)

    footer = document.sections[0].footer.paragraphs[0]
    footer.text = "Literature scan report generated from Markdown source"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("808080")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", default=str(SOURCE))
    parser.add_argument("--output", default=str(OUTPUT))
    args = parser.parse_args()
    build_docx(Path(args.source), Path(args.output))
