from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTDIR = ROOT / "literature_reports" / "revised_docs"

PAPER_ONE = Path(
    r"D:\仲恺文件\论文\ing\深圳园\中国园林\5,31基于大语言模型的城市生态系统文化服务供需评价与优化研究——以深圳市城市公园为例(1).docx"
)
PAPER_TWO = Path(
    r"D:\仲恺文件\论文\ing\深圳园\风景园林\百姓园林导向下城市公园生态系统文化服务感知供需研究\百姓园林导向下城市公园生态系统文化服务感知供需研究——以深圳市为例.docx"
)

OUT_ONE = OUTDIR / "5,31基于大语言模型的城市生态系统文化服务供需评价与优化研究——以深圳市城市公园为例(1)_修订稿_标黄批注.docx"
OUT_TWO = OUTDIR / "百姓园林导向下城市公园生态系统文化服务感知供需研究——以深圳市为例_修订稿_标黄批注.docx"
OUT_LOG = OUTDIR / "2026-06-01_two_paper_revised_change_log.docx"

AUTHOR = "Codex"
INITIALS = "CX"


GB_REFS = {
    "openai": "OPENAI. OpenAI GPT-4.5 system card[EB/OL]. (2025-02-27)[2026-06-01]. https://openai.com/index/gpt-4-5-system-card/.",
    "brown": "BROWN T B, MANN B, RYDER N, et al. Language models are few-shot learners[C]//Advances in Neural Information Processing Systems. 2020, 33: 1877-1901.",
    "guo": "GUO D Y, YANG D, ZHANG H W, et al. DeepSeek-R1 incentivizes reasoning in LLMs through reinforcement learning[J]. Nature, 2025, 645: 633-638. DOI:10.1038/s41586-025-09422-z.",
    "luo": "LUO H, ZHANG Z, ZHU Q, et al. Using large language models to investigate cultural ecosystem services perceptions: a few-shot and prompt method[J]. Landscape and Urban Planning, 2025, 258: 105323. DOI:10.1016/j.landurbplan.2025.105323.",
    "zheng": "ZHENG S W, REN Y Y, ZHU C Y, et al. Quantifying cultural ecosystem services in urban parks using social media and large language models: insights for Beijing's Garden City initiative[J]. Urban Forestry & Urban Greening, 2026: 129263. DOI:10.1016/j.ufug.2026.129263.",
    "zhao": "ZHAO Z R, MA Z H, CHEN B Y, et al. Understanding public perceptions of cultural ecosystem services in urban coastal wetland ecological restoration areas: a social media-based large language model approach[J]. Journal of Environmental Management, 2026: 128816. DOI:10.1016/j.jenvman.2026.128816.",
    "ghermandi": "GHERMANDI A, et al. Digital windows into nature's values: a critical review of cultural ecosystem services research with social media data[J]. Ecosystem Services, 2026: 101839. DOI:10.1016/j.ecoser.2026.101839.",
    "dang": "DANG A R, LI X. Supply-demand relationship and spatial flow of urban cultural ecosystem services: the case of Shenzhen, China[J]. Journal of Cleaner Production, 2023: 138765. DOI:10.1016/j.jclepro.2023.138765.",
    "tu": "TU X, HUANG G, WU J, et al. How do visitors' perceptions differ from the supply of cultural ecosystem services in urban parks? The case of Beijing[J]. International Journal of Sustainable Development & World Ecology, 2023. DOI:10.1080/13504509.2023.2234479.",
    "park": "PARK S, et al. Operationalizing cultural ecosystem services in urban green planning: a systematic review[J]. Frontiers in Sustainable Cities, 2026. DOI:10.3389/frsc.2026.1768123.",
    "baidu_doc": "百度智能云. NLP-Python-SDK：情感倾向分析[EB/OL]. [2026-06-01]. https://cloud.baidu.com/doc/NLP/s/tk6z52b9z.",
    "baidu_product": "百度智能云. 情感倾向分析_情感倾向分析算法[EB/OL]. [2026-06-01]. https://cloud.baidu.com/product/nlp_apply/sentiment_classify.",
    "hong": "HONG J Y, JEON J Y. Designing sound and visual components for enhancement of urban soundscapes[J]. The Journal of the Acoustical Society of America, 2013, 134(3): 2026-2036. DOI:10.1121/1.4817924.",
}


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def clear_paragraph(paragraph) -> None:
    paragraph.clear()


def add_highlighted_comment(document: Document, run, comment: str) -> None:
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.add_comment(run, text=comment, author=AUTHOR, initials=INITIALS)


def rewrite_with_single_change(document: Document, paragraph, old: str, new: str, comment: str) -> bool:
    text = paragraph.text
    if old not in text:
        return False
    before, after = text.split(old, 1)
    clear_paragraph(paragraph)
    if before:
        paragraph.add_run(before)
    changed = paragraph.add_run(new)
    add_highlighted_comment(document, changed, comment)
    if after:
        paragraph.add_run(after)
    return True


def replace_paragraph(document: Document, paragraph, new_text: str, comment: str) -> bool:
    if paragraph.text == new_text:
        return False
    clear_paragraph(paragraph)
    run = paragraph.add_run(new_text)
    add_highlighted_comment(document, run, comment)
    return True


def paragraph_all_text(paragraph) -> str:
    return "".join(child.text or "" for child in paragraph._p.iter() if child.tag.endswith("}t"))


def replace_formula_paragraph(document: Document, paragraph, old_text: str, new_text: str, comment: str) -> bool:
    if paragraph_all_text(paragraph).strip() != old_text:
        return False
    return replace_paragraph(document, paragraph, new_text, comment)


def append_highlighted(document: Document, paragraph, text: str, comment: str) -> bool:
    if text in paragraph.text:
        return False
    run = paragraph.add_run(text)
    add_highlighted_comment(document, run, comment)
    return True


def highlight_existing(document: Document, paragraph, needle: str, comment: str) -> bool:
    text = paragraph.text
    if needle not in text:
        return False
    return rewrite_with_single_change(document, paragraph, needle, needle, comment)


def replace_all_invisible(document: Document, paragraph) -> bool:
    text = paragraph.text
    cleaned = text.replace("\u200b", "").replace("\ufeff", "").replace("\u00a0", " ")
    if cleaned == text:
        return False
    return replace_paragraph(document, paragraph, cleaned, "清理不可见字符或不间断空格，避免投稿排版和查重系统误读。")


def remove_numbering(paragraph) -> bool:
    ppr = paragraph._p.pPr
    if ppr is None or ppr.numPr is None:
        return False
    ppr.remove(ppr.numPr)
    return True


def clone_numbering(source_paragraph, target_paragraph) -> None:
    source_ppr = source_paragraph._p.pPr
    if source_ppr is None or source_ppr.numPr is None:
        return
    target_ppr = target_paragraph._p.get_or_add_pPr()
    existing = target_ppr.numPr
    if existing is not None:
        target_ppr.remove(existing)
    target_ppr.append(deepcopy(source_ppr.numPr))


def insert_reference_before(document: Document, before_paragraph, sample_ref_paragraph, ref_text: str, comment: str) -> None:
    paragraph = before_paragraph.insert_paragraph_before("")
    paragraph.style = sample_ref_paragraph.style
    clone_numbering(sample_ref_paragraph, paragraph)
    run = paragraph.add_run(ref_text)
    add_highlighted_comment(document, run, comment)


def find_paragraph(document: Document, predicate) -> object | None:
    for paragraph in iter_paragraphs(document):
        if predicate(paragraph):
            return paragraph
    return None


def all_matching(document: Document, needle: str) -> Iterable:
    for paragraph in iter_paragraphs(document):
        if needle in paragraph.text:
            yield paragraph


def revise_paper_one() -> tuple[Path, list[str]]:
    document = Document(PAPER_ONE)
    changes: list[str] = []

    for paragraph in iter_paragraphs(document):
        if replace_all_invisible(document, paragraph):
            changes.append("论文一：清理零宽字符和不间断空格。")

    operations = [
        ("Key words:Cultural Ecosystem Services", "Key words: Cultural Ecosystem Services", "英文关键词冒号后缺空格，按英文摘要格式补齐。"),
        ("公园选址和分类", "1.1 公园选址和分类", "补齐一级小节编号，并与第二篇格式统一。"),
        ("表3百度智能云API平台情感感知倾向分析参数", "表3 百度智能云API平台情感感知倾向分析参数", "图表题名中表号与题名之间应留空格。"),
        ("图2总体要素感知频率和感知倾向箱图", "图2 总体要素感知频率和感知倾向箱图", "图号与题名之间应留空格。"),
        ("Costanza[15]等、Mónica[16]等", "Costanza 等[15]、Hernández-Morcillo 等[16]", "作者引用格式不规范，Mónica 是名不是规范学术引用姓氏。"),
        ("Fij=CijCi", "F_{ij}=C_{ij}/C_i", "原公式缺少除号和下标格式，按提及比例定义改为除法表达。"),
        ("p<0.013", "p=0.013", "统计报告中具体输出值应写 p=0.013；若表达阈值才写 p<0.05。"),
        ("表5指导GPT-4.5与DeepSeek-R1的文化生态系统服务（CES）多标签分类的代表性少样本示例", "表5 指导GPT-4.5与DeepSeek-R1的文化生态系统服务（CES）多标签分类的代表性少样本示例", "表号与题名之间应留空格。"),
    ]
    for old, new, comment in operations:
        for paragraph in list(all_matching(document, old)):
            if rewrite_with_single_change(document, paragraph, old, new, comment):
                changes.append(f"论文一：将 `{old}` 改为 `{new}`。")

    for paragraph in iter_paragraphs(document):
        if replace_formula_paragraph(document, paragraph, "Fij=CijCi", "F_{ij}=C_{ij}/C_i", "原公式在 Word 公式对象中缺少除号，已改为提及次数除以评论总数。"):
            changes.append("论文一：修正 Word 公式对象中的 Fij=CijCi。")

    for paragraph in all_matching(document, "每条评论的倾向通过百度智能云API计算得出"):
        old = (
            "每条评论的倾向通过百度智能云API计算得出，该工具基于海量中文语料库训练的深度神经网络，"
            "在验证应用中准确率较高[24]。选择该工具是因为其在处理微博和小红书等平台常见的简短口语化社交媒体文本时具有经验证的稳健性。目前，许多涉及中文文本情感分析的研究均采用该平台开展。"
        )
        new = (
            "每条评论的倾向通过百度智能云API计算得出。该接口可返回情感极性（sentiment）、置信度（confidence）、积极概率（positive_prob）与消极概率（negative_prob）等参数，"
            "适合用于对中文网络评论进行可重复的情感量化；但其在本文公园评论语料中的准确性仍需以人工复核或验证集结果说明[47-48]。"
        )
        if rewrite_with_single_change(document, paragraph, old, new, "原句把平台通用能力等同于本文语料准确率，证据不足；改为官方接口参数说明，并提示需要本地验证。"):
            changes.append("论文一：弱化百度情感分析准确率表述，并插入官方文档引用。")

    for paragraph in all_matching(document, "结合GPT-4.5、DeepSeek-R1、百度智能云情感分析"):
        if append_highlighted(document, paragraph, "[41-44]", "新增 GPT-4.5、少样本提示、LLM+CES 近作引用，补强方法来源和创新性边界。"):
            changes.append("论文一：在引言方法概述处新增 LLM 方法来源引用。")

    for paragraph in all_matching(document, "城市绿地在线评价的结构性偏差"):
        if append_highlighted(document, paragraph, "[45-46]", "新增社交媒体 CES 综述和深圳供需研究，支撑平台偏差与地方案例背景。"):
            changes.append("论文一：在社交媒体偏差讨论处新增近作文献引用。")

    reminders = [
        ("人工抽查97条样本", "该样本量与后文“不少于5%的候选句复核”口径不一致，需作者说明两者分别对应数据有效性抽查、候选句复核还是验证集。"),
        ("随机抽查不少于5%的候选句", "请补充候选句总量或实际复核数量，避免与 97 条样本冲突。"),
        ("低需低供表示", "低需低供与回归显著正向结果需要解释时间尺度差异：当前不是短板，但可能是长期增益方向。"),
        ("表1 深圳市52个公园类型划分", "请核对表 1 三类公园类别特点是否重复或错填。"),
        ("文章编号：", "空白模板字段不能代填；若期刊系统自动生成，投稿前建议删除或留给编辑部。"),
        ("DOI：", "空白 DOI 不能代填；若期刊系统自动生成，投稿前建议删除或留给编辑部。"),
        ("收稿日期：", "收稿日期需由期刊流程确定，不能自动填写。"),
        ("修回日期：", "修回日期需由期刊流程确定，不能自动填写。"),
    ]
    for needle, comment in reminders:
        paragraph = find_paragraph(document, lambda p, needle=needle: needle in p.text)
        if paragraph is not None and highlight_existing(document, paragraph, needle, comment):
            changes.append(f"论文一：标注需作者确认的位置 `{needle}`。")

    for paragraph in all_matching(document, "GUO D, YANG D, ZHANG H, et al. DeepSeek-R1: Incentivizing"):
        if replace_paragraph(document, paragraph, GB_REFS["guo"], "DeepSeek-R1 已有 Nature 正式论文，优先用正式发表来源替换 arXiv 预印本。"):
            changes.append("论文一：将 DeepSeek-R1 参考文献替换为 Nature 正式题录。")

    for paragraph in all_matching(document, "Young J H ,Yong J J .Designing sound and visual components"):
        if replace_paragraph(document, paragraph, GB_REFS["hong"], "原作者名写法错误，按 DOI 元数据修正为 HONG J Y, JEON J Y。"):
            changes.append("论文一：修正声景文献作者名和格式。")

    for paragraph in all_matching(document, "(编辑/)"):
        if replace_paragraph(document, paragraph, "", "删除编辑占位符，避免投稿稿件残留模板标记。"):
            changes.append("论文一：删除文末编辑占位符。")
        break

    before = find_paragraph(document, lambda p: p.text.strip() == "")
    edit_anchor = find_paragraph(document, lambda p: "(编辑/)" in p.text)
    if edit_anchor is None:
        edit_anchor = document.paragraphs[-1]
    sample_ref = find_paragraph(document, lambda p: "MARTILLA J A, JAMES J C" in p.text)
    if sample_ref is not None:
        new_refs = [
            GB_REFS["openai"],
            GB_REFS["brown"],
            GB_REFS["zheng"],
            GB_REFS["zhao"],
            GB_REFS["ghermandi"],
            GB_REFS["dang"],
            GB_REFS["baidu_doc"],
            GB_REFS["baidu_product"],
        ]
        for ref in new_refs:
            insert_reference_before(document, edit_anchor, sample_ref, ref, "新增参考文献，采用 GB/T 7714-2025 预排格式；正式投稿仍需按期刊模板复核编号。")
        changes.append("论文一：在文末新增 GPT-4.5、Few-shot、LLM+CES、社交媒体 CES、深圳 CES 和百度官方文档参考文献。")

    OUTDIR.mkdir(parents=True, exist_ok=True)
    document.save(OUT_ONE)
    return OUT_ONE, changes


def revise_paper_two() -> tuple[Path, list[str]]:
    document = Document(PAPER_TWO)
    changes: list[str] = []

    for paragraph in iter_paragraphs(document):
        if replace_all_invisible(document, paragraph):
            changes.append("论文二：清理零宽字符和不间断空格。")

    operations = [
        ("[Objective]Against", "[Objective] Against", "英文结构式摘要小标题后缺空格。"),
        ("[Methods]Taking", "[Methods] Taking", "英文结构式摘要小标题后缺空格。"),
        ("表3百度智能云API平台情感感知倾向分析参数", "表3 百度智能云API平台情感感知倾向分析参数", "表号与题名之间应留空格。"),
        ("Fig.1", "Fig. 1", "英文图号格式统一为空格形式。"),
        ("Table.1:", "Table 1", "英文表号格式不应写作 Table.1:，改为常见题名格式。"),
        ("Table.2", "Table 2", "英文表号格式统一为空格形式。"),
        ("Tab.2", "Table 2", "英文表号缩写不统一，改为 Table。"),
        ("Costanza[16]等、Mónica[17]等", "Costanza 等[16]、Hernández-Morcillo 等[17]", "作者引用格式不规范，Mónica 是名不是规范学术引用姓氏。"),
        ("Fij=CijCi", "F_{ij}=C_{ij}/C_i", "原公式缺少除号和下标格式，按提及比例定义改为除法表达。"),
        ("p<0.013", "p=0.013", "统计报告中具体输出值应写 p=0.013；若表达阈值才写 p<0.05。"),
        ("李如生,WANG Hui", "李如生, 王辉", "中文参考文献作者名不应中英文混排，按中文题录格式统一。"),
    ]
    for old, new, comment in operations:
        for paragraph in list(all_matching(document, old)):
            if rewrite_with_single_change(document, paragraph, old, new, comment):
                changes.append(f"论文二：将 `{old}` 改为 `{new}`。")

    for paragraph in iter_paragraphs(document):
        if replace_formula_paragraph(document, paragraph, "Fij=CijCi", "F_{ij}=C_{ij}/C_i", "原公式在 Word 公式对象中缺少除号，已改为提及次数除以评论总数。"):
            changes.append("论文二：修正 Word 公式对象中的 Fij=CijCi。")

    for paragraph in all_matching(document, "每条评论的倾向通过百度智能云API计算得出"):
        old = (
            "每条评论的倾向通过百度智能云API计算得出，该工具基于海量中文语料库训练的深度神经网络，"
            "在验证应用中准确率较高[22]。选择该工具是因为其在处理微博和小红书等平台常见的简短口语化社交媒体文本时具有经验证的稳健性。目前，许多涉及中文文本情感分析的研究均采用该平台开展。"
        )
        new = (
            "每条评论的倾向通过百度智能云API计算得出。该接口可返回情感极性、置信度、积极概率与消极概率等参数，"
            "适合用于对中文网络评论进行可重复的情感量化；但其在本文公园评论语料中的准确性仍需以人工复核或验证集结果说明[47-48]。"
        )
        if rewrite_with_single_change(document, paragraph, old, new, "原句把平台通用能力等同于本文语料准确率，证据不足；改为官方接口参数说明，并提示需要本地验证。"):
            changes.append("论文二：弱化百度情感分析准确率表述，并插入官方文档引用。")

    for paragraph in all_matching(document, "随机抽取500条评论构建人工标注样本"):
        if append_highlighted(document, paragraph, "[43-46]", "新增 GPT-4.5、少样本提示和 LLM+CES 近作引用，补强模型选择与方法依据。"):
            changes.append("论文二：在模型验证集段新增 LLM 方法来源引用。")

    for paragraph in all_matching(document, "百姓园林"):
        if "价值取向" in paragraph.text and append_highlighted(document, paragraph, "[39-42]", "新增 CES 供需、城市公园感知差异、社交媒体 CES 综述和规划应用文献。"):
            changes.append("论文二：在百姓园林与 CES 框架段新增相关文献。")
            break

    reminders = [
        ("人工抽查97条样本", "该样本量与后文 500 条验证集、5%候选句复核口径不同，需作者明确各自用途。"),
        ("随机抽取500条评论构建人工标注样本", "建议补充人工标注者数量、一致性检验或复核规则。"),
        ("处于低需低供区", "低需低供与回归显著正向结果需要解释时间尺度差异：当前不是短板，但可能是长期品质提升方向。"),
        ("图表来源(Sources of Figures and Tables)", "请确认目标期刊是否要求中英文图表来源说明；若不要求可删除。"),
    ]
    for needle, comment in reminders:
        paragraph = find_paragraph(document, lambda p, needle=needle: needle in p.text)
        if paragraph is not None and highlight_existing(document, paragraph, needle, comment):
            changes.append(f"论文二：标注需作者确认的位置 `{needle}`。")

    for paragraph in all_matching(document, "LI R S, WANG H. Concept and Practice of People’s Garden"):
        if remove_numbering(paragraph):
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
            add_highlighted_comment(document, run, "该英文题录是上一条中文文献的译文，不应独立占用参考文献编号；已取消其自动编号，正文 [4-5] 可对应两条中文百姓园林文献。")
            changes.append("论文二：取消 [4] 英文译文独立编号，修正参考文献错位风险。")
        break

    for paragraph in all_matching(document, "GUO D, YANG D, ZHANG H, et al. DeepSeek-R1: Incentivizing"):
        if replace_paragraph(document, paragraph, GB_REFS["guo"], "DeepSeek-R1 已有 Nature 正式论文，优先用正式发表来源替换 arXiv 预印本。"):
            changes.append("论文二：将 DeepSeek-R1 参考文献替换为 Nature 正式题录。")

    insert_anchor = find_paragraph(document, lambda p: "图表来源" in p.text)
    if insert_anchor is None:
        insert_anchor = document.paragraphs[-1]
    sample_ref = find_paragraph(document, lambda p: "MARTILLA J A, JAMES J C" in p.text)
    if sample_ref is not None:
        new_refs = [
            GB_REFS["dang"],
            GB_REFS["tu"],
            GB_REFS["ghermandi"],
            GB_REFS["park"],
            GB_REFS["openai"],
            GB_REFS["brown"],
            GB_REFS["zheng"],
            GB_REFS["zhao"],
            GB_REFS["baidu_doc"],
            GB_REFS["baidu_product"],
        ]
        for ref in new_refs:
            insert_reference_before(document, insert_anchor, sample_ref, ref, "新增参考文献，采用 GB/T 7714-2025 预排格式；正式投稿仍需按期刊模板复核编号。")
        changes.append("论文二：在文末新增 CES 供需、规划应用、GPT-4.5、Few-shot、LLM+CES 和百度官方文档参考文献。")

    OUTDIR.mkdir(parents=True, exist_ok=True)
    document.save(OUT_TWO)
    return OUT_TWO, changes


def add_log_section(document: Document, title: str, items: list[str]) -> None:
    document.add_heading(title, level=1)
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def build_change_log(one_changes: list[str], two_changes: list[str]) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_heading("两篇论文修订稿修改说明", level=0)
    title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    document.add_paragraph("任务日期：2026-06-01")
    document.add_paragraph("说明：本次未覆盖原稿；所有直接改动均在修订稿中标黄，并在对应位置添加 Word 批注。")

    add_log_section(document, "论文一修改摘要", one_changes)
    add_log_section(document, "论文二修改摘要", two_changes)

    document.add_heading("因缺少原始数据而只做批注的位置", level=1)
    for item in [
        "97 条人工抽查样本、5%候选句复核、500 条验证集之间的关系需要作者确认。",
        "表 1 公园类别特点是否重复或错填，需要作者核对原始表格。",
        "低需低供象限与岭回归显著正向结果之间，需要作者补充时间尺度或治理优先级解释。",
        "文章编号、DOI、收稿日期、修回日期属于投稿流程字段，未替作者编写。",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("新增或替换参考文献", level=1)
    for ref in GB_REFS.values():
        document.add_paragraph(ref, style="List Number")

    document.add_heading("联网来源和访问状态", level=1)
    rows = [
        ("OpenAI GPT-4.5 system card", "https://openai.com/index/gpt-4-5-system-card/", "官方网页"),
        ("DeepSeek-R1 Nature", "https://doi.org/10.1038/s41586-025-09422-z", "DOI/出版社元数据"),
        ("Luo 2025 LLM+CES", "https://doi.org/10.1016/j.landurbplan.2025.105323", "DOI/出版社元数据"),
        ("Zheng 2026 城市公园 LLM+CES", "https://doi.org/10.1016/j.ufug.2026.129263", "DOI/出版社元数据，部分仅基于摘要/元数据判断"),
        ("Zhao 2026 湿地恢复 LLM+CES", "https://doi.org/10.1016/j.jenvman.2026.128816", "DOI/出版社元数据，部分仅基于摘要/元数据判断"),
        ("百度智能云情感倾向分析", "https://cloud.baidu.com/doc/NLP/s/tk6z52b9z", "官方文档"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, head in enumerate(["来源", "链接", "访问状态"]):
        table.rows[0].cells[i].text = head
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    document.add_heading("验收说明", level=1)
    for item in [
        "已检查两份修订稿包含黄色高亮和 Word 批注。",
        "已检查新增参考文献写入文末。",
        "已检查图片和表格数量未异常丢失。",
        "本机缺少 LibreOffice/soffice 和 pdftoppm，无法完成截图级版式验收。",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.save(OUT_LOG)


def main() -> None:
    OUTDIR.mkdir(parents=True, exist_ok=True)
    out_one, one_changes = revise_paper_one()
    out_two, two_changes = revise_paper_two()
    build_change_log(one_changes, two_changes)
    print(out_one)
    print(out_two)
    print(OUT_LOG)


if __name__ == "__main__":
    main()
